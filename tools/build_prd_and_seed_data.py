from __future__ import annotations

import csv
import html
import os
import re
import textwrap
from datetime import date
from pathlib import Path
from urllib.error import HTTPError, URLError
from urllib.parse import urljoin
from urllib.request import Request, urlopen

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = ROOT / "data"
OUTPUT_DIR = ROOT / "outputs"
REPORT_PATH = OUTPUT_DIR / "乒器通微信小程序产品需求报告.docx"
TODAY = "2026-07-05"

BUTTERFLY_BASE = "https://www.butterfly-global.com"
BUTTERFLY_ATHLETE_URLS = [
    "https://www.butterfly-global.com/en/sponsoring/detail/fan-zhendong.html",
    "https://www.butterfly-global.com/en/sponsoring/detail/harimoto-tomokazu.html",
    "https://www.butterfly-global.com/en/sponsoring/detail/harimoto-miwa.html",
    "https://www.butterfly-global.com/en/sponsoring/detail/lin-shidong.html",
    "https://www.butterfly-global.com/en/sponsoring/detail/zhu-yuling.html",
    "https://www.butterfly-global.com/en/sponsoring/detail/lin-yunju.html",
    "https://www.butterfly-global.com/en/sponsoring/detail/matsushima-sora.html",
    "https://www.butterfly-global.com/en/sponsoring/detail/qiu-dang.html",
    "https://www.butterfly-global.com/en/sponsoring/detail/togami-shunsuke.html",
    "https://www.butterfly-global.com/en/sponsoring/detail/franziska-patrick.html",
    "https://www.butterfly-global.com/en/sponsoring/detail/ovtcharov-dimitrij.html",
    "https://www.butterfly-global.com/en/sponsoring/detail/lee-sangsu.html",
    "https://www.butterfly-global.com/en/sponsoring/detail/jha-kanak.html",
    "https://www.butterfly-global.com/en/sponsoring/detail/assar-omar.html",
    "https://www.butterfly-global.com/en/sponsoring/detail/cheng-i-ching.html",
    "https://www.butterfly-global.com/en/sponsoring/detail/diaz-adriana.html",
]

FALLBACK_BUTTERFLY_ROWS = [
    ("Fan Zhendong", "CHN", "July 19, 2022", "Fan Zhendong ALC", "blade", "https://www.butterfly-global.com/en/products/detail/37221.html", "https://www.butterfly-global.com/en/sponsoring/detail/fan-zhendong.html"),
    ("Fan Zhendong", "CHN", "July 19, 2022", "Dignics 09C", "rubber", "https://www.butterfly-global.com/en/products/detail/06070.html", "https://www.butterfly-global.com/en/sponsoring/detail/fan-zhendong.html"),
    ("Tomokazu Harimoto", "JPN", "October 11, 2025", "Harimoto Tomokazu Innerforce Super ALC", "blade", "https://www.butterfly-global.com/en/products/detail/37331.html", "https://www.butterfly-global.com/en/sponsoring/detail/harimoto-tomokazu.html"),
    ("Tomokazu Harimoto", "JPN", "October 11, 2025", "Zyre 03", "rubber", "https://www.butterfly-global.com/en/products/detail/06140.html", "https://www.butterfly-global.com/en/sponsoring/detail/harimoto-tomokazu.html"),
    ("Miwa Harimoto", "JPN", "December 18, 2025", "Harimoto Tomokazu Innerforce Super ALC", "blade", "https://www.butterfly-global.com/en/products/detail/37331.html", "https://www.butterfly-global.com/en/sponsoring/detail/harimoto-miwa.html"),
    ("Miwa Harimoto", "JPN", "December 18, 2025", "Dignics 09C", "rubber", "https://www.butterfly-global.com/en/products/detail/06070.html", "https://www.butterfly-global.com/en/sponsoring/detail/harimoto-miwa.html"),
    ("Miwa Harimoto", "JPN", "December 18, 2025", "Dignics 05", "rubber", "https://www.butterfly-global.com/en/products/detail/06040.html", "https://www.butterfly-global.com/en/sponsoring/detail/harimoto-miwa.html"),
    ("Lin Shidong", "CHN", "June 1, 2022", "Dignics 09C", "rubber", "https://www.butterfly-global.com/en/products/detail/06070.html", "https://www.butterfly-global.com/en/sponsoring/detail/lin-shidong.html"),
    ("Zhu Yuling", "MAC", "February 16, 2026", "Tenergy 05 Hard", "rubber", "https://www.butterfly-global.com/en/products/detail/06030.html", "https://www.butterfly-global.com/en/sponsoring/detail/zhu-yuling.html"),
    ("Zhu Yuling", "MAC", "February 16, 2026", "Dignics 09C", "rubber", "https://www.butterfly-global.com/en/products/detail/06070.html", "https://www.butterfly-global.com/en/sponsoring/detail/zhu-yuling.html"),
    ("Lin Yun-Ju", "TPE", "January 5, 2026", "Viscaria Super ALC", "blade", "https://www.butterfly-global.com/en/products/detail/37191.html", "https://www.butterfly-global.com/en/sponsoring/detail/lin-yunju.html"),
    ("Lin Yun-Ju", "TPE", "January 5, 2026", "Zyre 03", "rubber", "https://www.butterfly-global.com/en/products/detail/06140.html", "https://www.butterfly-global.com/en/sponsoring/detail/lin-yunju.html"),
    ("Lin Yun-Ju", "TPE", "January 5, 2026", "Dignics 05", "rubber", "https://www.butterfly-global.com/en/products/detail/06040.html", "https://www.butterfly-global.com/en/sponsoring/detail/lin-yunju.html"),
    ("Shunsuke Togami", "JPN", "May 21, 2026", "Fan Zhendong ALC", "blade", "https://www.butterfly-global.com/en/products/detail/37221.html", "https://www.butterfly-global.com/en/sponsoring/detail/togami-shunsuke.html"),
    ("Patrick Franziska", "GER", "June 15, 2026", "Franziska Innerforce ZLC", "blade", "https://www.butterfly-global.com/en/products/detail/37121.html", "https://www.butterfly-global.com/en/sponsoring/detail/franziska-patrick.html"),
    ("Kanak Jha", "USA", "February 12, 2026", "Timo Boll ALC", "blade", "https://www.butterfly-global.com/en/products/detail/35861.html", "https://www.butterfly-global.com/en/sponsoring/detail/jha-kanak.html"),
]


BRANDS = [
    ("DHS / 红双喜", "中国", "底板、套胶、成品拍、球台", "国货核心品牌；狂飙系列是中国市场高频选项", "https://www.dhs-sports.com/"),
    ("Butterfly / 蝴蝶", "日本", "底板、套胶、服装鞋包", "国际高端主流；Viscaria、Tenergy、Dignics 系列认知度高", "https://www.butterfly-global.com/en/"),
    ("STIGA / 斯帝卡", "瑞典", "底板、套胶、成品拍", "Clipper、Carbonado、DNA 系列常见", "https://www.stiga.com/"),
    ("Nittaku / 尼塔库", "日本", "底板、套胶、球", "Acoustic、Fastarc G-1 等常见", "https://www.nittaku.com/"),
    ("Yasaka / 亚萨卡", "日本/瑞典", "底板、套胶", "Ma Lin Extra Offensive、Rakza 系列常见", "https://www.yasaka.se/"),
    ("Donic / 多尼克", "德国", "底板、套胶", "Waldner、Bluefire、Bluestorm 系列常见", "https://www.donic.com/"),
    ("TIBHAR / 挺拔", "德国", "底板、套胶", "Evolution、Hybrid K3、Lebrun 系列常见", "https://www.tibhar.com/"),
    ("XIOM / 骄猛", "韩国", "底板、套胶", "Vega、Omega、Jekyll & Hyde 系列常见", "https://www.xiom.global/"),
    ("VICTAS", "日本", "底板、套胶、颗粒胶", "V>15、V>20、Curl、防守板线常见", "https://www.victas.com/"),
    ("Andro / 岸度", "德国", "底板、套胶", "Rasanter、Hexer 系列常见", "https://www.andro.de/"),
    ("JOOLA / 优拉", "德国/美国", "底板、套胶、赛事器材", "Dynaryz、Rhyzen、Vyzaryz 系列常见", "https://joola.com/"),
    ("Yinhe / 银河", "中国", "底板、套胶", "入门与性价比市场常见", "https://www.galaxysport.com/"),
    ("729 / 友谊", "中国", "套胶、底板", "Battle、Focus、729 系列常见", "https://www.729sports.com/"),
    ("Sanwei / 三维", "中国", "底板、套胶", "Target、Fextra 等性价比产品常见", "https://www.sanweisport.com/"),
    ("Palio / 拍里奥", "中国", "套胶、成品拍", "AK47、CJ8000 等入门市场常见", "https://www.palioett.com/"),
    ("Double Fish / 双鱼", "中国", "球台、球、套胶、成品拍", "大众市场和赛事器材常见", "https://www.doublefish.com/"),
]


MANUAL_EQUIPMENT = [
    ("DHS / 红双喜", "Hurricane Long 5 / 狂飙龙5", "blade", "高阶弧圈快攻", "高", "需品牌官网或授权店参数校验"),
    ("DHS / 红双喜", "301X", "blade", "进阶弧圈快攻", "中高", "需品牌官网或授权店参数校验"),
    ("DHS / 红双喜", "Power G7 / 劲极7", "blade", "入门到进阶快攻弧圈", "中", "需品牌官网或授权店参数校验"),
    ("Butterfly / 蝴蝶", "Viscaria", "blade", "两面弧圈快攻", "高", "后续补充官方产品页"),
    ("Butterfly / 蝴蝶", "Fan Zhendong ALC", "blade", "两面弧圈快攻", "高", "已可由 Butterfly 产品页校验"),
    ("Butterfly / 蝴蝶", "Timo Boll ALC", "blade", "快攻弧圈", "高", "已可由 Butterfly 产品页校验"),
    ("STIGA / 斯帝卡", "Clipper Wood", "blade", "近台快攻弧圈", "中高", "需品牌官网或授权店参数校验"),
    ("Yasaka / 亚萨卡", "Ma Lin Extra Offensive", "blade", "弧圈快攻", "中", "需品牌官网或授权店参数校验"),
    ("Nittaku / 尼塔库", "Acoustic", "blade", "控制弧圈", "中高", "需品牌官网或授权店参数校验"),
    ("VICTAS", "Koji Matsushita", "blade", "削球防守", "中高", "需品牌官网或授权店参数校验"),
    ("DHS / 红双喜", "Hurricane 3 Neo / 狂飙3 NEO", "rubber", "正手粘性弧圈", "中高", "需品牌官网或授权店参数校验"),
    ("DHS / 红双喜", "Skyline 2 / 天极2", "rubber", "正手粘性弧圈", "中高", "需品牌官网或授权店参数校验"),
    ("Butterfly / 蝴蝶", "Dignics 09C", "rubber", "高摩擦粘弹两用", "高", "已可由 Butterfly 产品页校验"),
    ("Butterfly / 蝴蝶", "Dignics 05", "rubber", "高旋转高弹弧圈", "高", "已可由 Butterfly 产品页校验"),
    ("Butterfly / 蝴蝶", "Tenergy 05", "rubber", "高弹弧圈", "高", "后续补充官方产品页"),
    ("Butterfly / 蝴蝶", "Rozena", "rubber", "进阶反手/均衡", "中", "后续补充官方产品页"),
    ("Yasaka / 亚萨卡", "Rakza 7", "rubber", "均衡弧圈", "中", "需品牌官网或授权店参数校验"),
    ("Nittaku / 尼塔库", "Fastarc G-1", "rubber", "反手/正手高旋转", "中高", "需品牌官网或授权店参数校验"),
    ("TIBHAR / 挺拔", "Evolution MX-P", "rubber", "高弹进攻", "中高", "需品牌官网或授权店参数校验"),
    ("XIOM / 骄猛", "Vega Europe", "rubber", "入门进阶反手", "中", "需品牌官网或授权店参数校验"),
    ("VICTAS", "V>15 Extra", "rubber", "速度旋转型反手/正手", "中高", "需品牌官网或授权店参数校验"),
    ("729 / 友谊", "Battle 2 / 奔腾2", "rubber", "性价比粘性正手", "中", "需品牌官网或授权店参数校验"),
    ("Yinhe / 银河", "Mercury II / 水星2", "rubber", "入门训练", "低", "需品牌官网或授权店参数校验"),
    ("Palio / 拍里奥", "AK47", "rubber", "性价比反手", "低中", "需品牌官网或授权店参数校验"),
]


COMMON_COMBOS = [
    ("入门练基本功", "300-600", "银河 N11S / 红双喜 劲极7", "水星2 / 729 普及型粘性套胶", "水星2 / AK47", "初学者、动作未定型", "C", "社区常见搭配种子，需淘宝/小红书授权采样验证"),
    ("进阶控制弧圈", "600-1200", "Yasaka Ma Lin Extra Offensive", "Rakza 7 / Hurricane 3 Neo", "Rakza 7 / Vega Europe", "一年以上训练、追求稳定上台", "C", "社区常见搭配种子，需授权采样验证"),
    ("国套正手 + 外套反手", "800-1600", "DHS 301X / Fang Bo B2X", "Hurricane 3 Neo", "Fastarc G-1 / Rakza 7 / Rozena", "中级到高级两面弧圈", "C", "社区常见搭配种子，需授权采样验证"),
    ("经典外置 ALC", "1600-3000", "Butterfly Viscaria / Fan Zhendong ALC", "Hurricane 3 Neo / Dignics 09C", "Dignics 05 / Tenergy 05 / Rozena", "高级快弧、发力较完整", "B", "Butterfly 产品可校验；具体搭配需平台采样验证"),
    ("内置持球弧圈", "1200-2600", "Butterfly Innerforce ALC / DHS Long 5", "Hurricane 3 Neo / Dignics 09C", "Dignics 05 / Fastarc G-1", "中远台弧圈、重视持球", "C", "社区常见搭配种子，需授权采样验证"),
    ("近台快攻压迫", "900-1800", "STIGA Clipper / Clipper CR", "Hurricane 3 Neo / DNA Platinum", "DNA Pro / V>15 Extra", "近台快攻、借力能力强", "C", "社区常见搭配种子，需授权采样验证"),
    ("削球/长胶防守", "700-1800", "VICTAS Koji Matsushita / 防守型底板", "粘性正胶/反胶", "Curl P1V / 长胶", "削球、防守反击、特殊打法", "C", "需专项社区与教练验证"),
    ("儿童轻量训练", "200-700", "轻量五夹木底板", "中软控制型套胶", "中软控制型套胶", "7-12 岁、控重优先", "C", "需教练和家长反馈验证"),
]


QUESTIONS = [
    (1, "10:8 领先且你有发球权，你更想怎么处理？", "A. 发短下旋抢第三板;B. 发奔球偷袭;C. 先保证不失误;D. 发对方别扭位等他先急", "risk,serve,tempo"),
    (2, "反手被连续压住时，你第一反应是什么？", "A. 拧起来抢主动;B. 快撕快带顶回去;C. 先挡稳找落点;D. 退半步拉一板质量", "backhand,risk,defense"),
    (3, "你最享受赢球的方式是？", "A. 发抢一套带走;B. 多拍相持磨到对方崩;C. 变化旋转骗失误;D. 防住以后反击", "tempo,spin,defense"),
    (4, "新球拍到手第一天，你会先测什么？", "A. 爆冲上限;B. 小球和台内控制;C. 反手借力;D. 连续拉的稳定性", "equipment_sense"),
    (5, "比分胶着时，你会更愿意？", "A. 坚持自己的强项;B. 临时变化套路;C. 降低失误率;D. 看对手表情调整", "mental,tactics"),
    (6, "对手搓得很转很低，你会？", "A. 强行拧/挑;B. 劈长到底线;C. 摆短等机会;D. 拉高吊先上台", "receive,spin,risk"),
    (7, "你选择正手胶皮最看重？", "A. 旋转和咬球;B. 速度和通透;C. 容错和控制;D. 不吃发力、轻松出质量", "rubber_preference"),
    (8, "你选择底板最怕？", "A. 太肉打不穿;B. 太弹小球飞;C. 太重手腕累;D. 手感太虚不踏实", "blade_preference"),
    (9, "一周训练频率是？", "A. 4 次以上;B. 2-3 次;C. 1 次左右;D. 偶尔打打", "training_frequency"),
    (10, "你更像哪种站位？", "A. 贴台压迫;B. 中近台相持;C. 中远台拉冲;D. 退台削防", "distance,style"),
    (11, "看到对方反手弱，你会？", "A. 一直打到他崩;B. 打几板后突然变正手;C. 不刻意针对，先打稳;D. 用旋转和节奏折磨", "tactics"),
    (12, "你更讨厌哪类失误？", "A. 发力出界;B. 吃发球;C. 连续相持掉链子;D. 错过主动进攻机会", "pain_point"),
    (13, "你对器材重量的态度？", "A. 重一点也行，质量优先;B. 必须轻快;C. 看手感平衡;D. 肩肘不舒服，要减负", "weight,injury"),
    (14, "别人说你打球像什么？", "A. 节奏快，压迫强;B. 旋转多，线路贼;C. 稳，不容易送;D. 很会耗，很难打死", "persona"),
    (15, "如果预算翻倍，你最想升级？", "A. 底板上限;B. 正手胶皮质量;C. 反手稳定性;D. 先请教练再换装备", "budget_strategy"),
]


CN_ATHLETE_WATCHLIST = [
    ("樊振东", "CHN", "已入 A 级证据：Butterfly 官方赞助页", "Butterfly 官方页显示 Fan Zhendong ALC、Dignics 09C，日期为 2022-07-19；需继续补充不同阶段装备变化"),
    ("林诗栋", "CHN", "已入 A 级证据但信息不完整", "Butterfly 官方页显示 Dignics 09C，日期为 2022-06-01；底板与另一面套胶需补证"),
    ("马龙", "CHN", "待验证", "优先查 DHS 官方、赞助物料、比赛高清图和权威采访；不得直接采用论坛传闻"),
    ("王楚钦", "CHN", "待验证", "优先查 DHS 官方、赞助物料、比赛高清图和权威采访；记录不同年份变化"),
    ("梁靖崑", "CHN", "待验证", "优先查品牌官方、国家队公开物料、赛事图像"),
    ("林高远", "CHN", "待验证", "优先查品牌官方、国家队公开物料、赛事图像"),
    ("孙颖莎", "CHN", "待验证", "优先查 DHS 官方、赞助物料、比赛高清图和权威采访；不得直接采用论坛传闻"),
    ("王曼昱", "CHN", "待验证", "优先查品牌官方、国家队公开物料、赛事图像"),
    ("陈梦", "CHN", "待验证", "优先查品牌官方、国家队公开物料、赛事图像"),
    ("王艺迪", "CHN", "待验证", "优先查品牌官方、国家队公开物料、赛事图像"),
    ("陈幸同", "CHN", "待验证", "优先查品牌官方、国家队公开物料、赛事图像"),
    ("钱天一", "CHN", "待验证", "优先查品牌官方、国家队公开物料、赛事图像"),
    ("蒯曼", "CHN", "待验证", "优先查品牌官方、国家队公开物料、赛事图像"),
    ("向鹏", "CHN", "待验证", "优先查品牌官方、国家队公开物料、赛事图像"),
]


SOURCES = [
    ("Butterfly Sponsoring", "品牌官方赞助运动员与装备页", "https://www.butterfly-global.com/en/sponsoring/", "运动员装备 A 级证据"),
    ("Butterfly Products", "品牌官方产品参数页", "https://www.butterfly-global.com/en/products/", "底板/套胶参数 A 级证据"),
    ("WTT Rankings", "世界排名公开页面", "https://www.worldtabletennis.com/rankings", "用于识别顶尖运动员范围，不直接证明装备"),
    ("微信 code2Session", "微信小程序登录服务端接口", "https://developers.weixin.qq.com/miniprogram/dev/OpenApiDoc/user-login/code2Session.html", "微信登录与账号体系"),
    ("微信 msgSecCheck", "微信内容安全文本审核接口", "https://developers.weixin.qq.com/miniprogram/dev/OpenApiDoc/sec-center/sec-check/msgSecCheck.html", "评论/帖子审核"),
    ("微信 ad 组件", "微信小程序广告组件", "https://developers.weixin.qq.com/miniprogram/dev/component/ad.html", "广告位预留"),
    ("淘宝开放平台", "淘宝/天猫商品数据授权入口", "https://open.taobao.com/", "商品/价格/销量类数据必须走授权接口或人工样本"),
    ("小红书蒲公英", "小红书商业合作平台", "https://pgy.xiaohongshu.com/", "达人/内容合作数据优先走授权或人工合规样本"),
]


def fetch_text(url: str, timeout: int = 15) -> str | None:
    try:
        req = Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urlopen(req, timeout=timeout) as response:
            return response.read().decode("utf-8", "replace")
    except (HTTPError, URLError, TimeoutError, OSError):
        return None


def clean_html(raw: str) -> str:
    text = re.sub(r"<br\s*/?>", "\n", raw, flags=re.I)
    text = re.sub(r"<.*?>", " ", text)
    text = html.unescape(text)
    return re.sub(r"\s+", " ", text).strip()


def infer_category(product_name: str, product_url: str) -> str:
    lower = product_name.lower()
    if any(k in lower for k in ["dignics", "tenergy", "zyre", "rozena", "bryce", "spinart"]):
        return "rubber"
    if "/products/detail/06" in product_url:
        return "rubber"
    return "blade"


def parse_country(display_name: str) -> tuple[str, str]:
    match = re.search(r"^(.*?)\s*\(([A-Z]{3})\)", display_name)
    if match:
        return match.group(1).strip(), match.group(2).strip()
    return display_name.strip(), ""


def collect_butterfly_athlete_rows() -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    for url in BUTTERFLY_ATHLETE_URLS:
        page = fetch_text(url, timeout=15)
        if not page:
            continue
        title_match = re.search(r"<title>(.*?)</title>", page, flags=re.S)
        display_name = clean_html(title_match.group(1)).split("|")[0].strip() if title_match else ""
        athlete_name, country = parse_country(display_name)

        block = ""
        start = page.find("playersDetail_product")
        if start >= 0:
            end = page.find("</section>", start)
            block = page[start:end]
        if not block:
            continue

        date_match = re.search(r"as of ([^<]+)", block)
        source_date = clean_html(date_match.group(1)) if date_match else ""
        product_matches = re.finditer(
            r'<a href="([^"]+)" class="imgSet">\s*<img[^>]+alt="([^"]*)"[^>]*>\s*<div class="imgSet_caption">(.*?)</div>',
            block,
            flags=re.S,
        )
        seq = 0
        for match in product_matches:
            product_name = clean_html(match.group(3))
            if not product_name:
                continue
            product_url = urljoin(BUTTERFLY_BASE, match.group(1))
            category = infer_category(product_name, product_url)
            seq += 1
            rows.append(
                {
                    "athlete_name": athlete_name,
                    "country": country,
                    "brand": "Butterfly / 蝴蝶",
                    "equipment_name": product_name,
                    "equipment_category": category,
                    "slot_sequence": str(seq),
                    "source_date": source_date,
                    "source_url": url,
                    "product_url": product_url,
                    "evidence_level": "A",
                    "evidence_type": "official_brand_sponsoring_page",
                    "notes": "官网页面给出 Equipment as of 日期；同名套胶出现两次时保留顺序，后续可人工标注正手/反手。",
                }
            )
    if rows:
        return rows

    fallback = []
    for seq, (athlete, country, source_date, equipment, category, product_url, source_url) in enumerate(FALLBACK_BUTTERFLY_ROWS, 1):
        fallback.append(
            {
                "athlete_name": athlete,
                "country": country,
                "brand": "Butterfly / 蝴蝶",
                "equipment_name": equipment,
                "equipment_category": category,
                "slot_sequence": str(seq),
                "source_date": source_date,
                "source_url": source_url,
                "product_url": product_url,
                "evidence_level": "A",
                "evidence_type": "official_brand_sponsoring_page",
                "notes": "网络采集失败时使用已人工核验的备用种子记录。",
            }
        )
    return fallback


def parse_product_specs(product_name: str, product_url: str) -> dict[str, str]:
    page = fetch_text(product_url, timeout=15)
    specs = {
        "brand": "Butterfly / 蝴蝶",
        "name": product_name,
        "category": infer_category(product_name, product_url),
        "product_url": product_url,
        "type": "",
        "speed": "",
        "spin": "",
        "arc": "",
        "sponge_hardness": "",
        "reaction_property": "",
        "vibration_property": "",
        "structure": "",
        "thickness": "",
        "source_status": "not_fetched",
    }
    if not page:
        return specs
    specs["source_status"] = "fetched"
    for th, td in re.findall(r"<tr>\s*<th>(.*?)</th>\s*<td>(.*?)</td>\s*</tr>", page, flags=re.S):
        key = clean_html(th).lower()
        value = clean_html(td)
        if key == "type":
            specs["type"] = value
        elif key == "speed":
            specs["speed"] = value
        elif key == "spin":
            specs["spin"] = value
        elif key == "arc":
            specs["arc"] = value
        elif key == "sponge hardness":
            specs["sponge_hardness"] = value
        elif key == "reaction property":
            specs["reaction_property"] = value
        elif key == "vibration property":
            specs["vibration_property"] = value
        elif key == "structure":
            specs["structure"] = value
        elif key == "thickness":
            specs["thickness"] = value
    return specs


def write_csv(path: Path, rows: list[dict[str, str]], fieldnames: list[str]) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        for row in rows:
            writer.writerow({k: row.get(k, "") for k in fieldnames})


def write_seed_data() -> dict[str, Path]:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    athlete_rows = collect_butterfly_athlete_rows()
    write_csv(
        DATA_DIR / "athlete_equipment_evidence_seed.csv",
        athlete_rows,
        [
            "athlete_name",
            "country",
            "brand",
            "equipment_name",
            "equipment_category",
            "slot_sequence",
            "source_date",
            "source_url",
            "product_url",
            "evidence_level",
            "evidence_type",
            "notes",
        ],
    )

    unique_products = {}
    for row in athlete_rows:
        unique_products[row["product_url"]] = row["equipment_name"]
    product_rows = [parse_product_specs(name, url) for url, name in unique_products.items()]
    write_csv(
        DATA_DIR / "butterfly_product_specs_seed.csv",
        product_rows,
        [
            "brand",
            "name",
            "category",
            "product_url",
            "type",
            "speed",
            "spin",
            "arc",
            "sponge_hardness",
            "reaction_property",
            "vibration_property",
            "structure",
            "thickness",
            "source_status",
        ],
    )

    brand_rows = [
        {
            "brand": b,
            "origin": o,
            "main_categories": c,
            "market_note": n,
            "official_url": u,
            "verification_status": "source_url_seed",
        }
        for b, o, c, n, u in BRANDS
    ]
    write_csv(
        DATA_DIR / "brand_catalog_seed.csv",
        brand_rows,
        ["brand", "origin", "main_categories", "market_note", "official_url", "verification_status"],
    )

    equipment_rows = [
        {
            "brand": b,
            "name": n,
            "category": c,
            "recommended_style": s,
            "price_band": p,
            "data_status": "manual_seed_needs_spec_verification",
            "notes": note,
        }
        for b, n, c, s, p, note in MANUAL_EQUIPMENT
    ]
    write_csv(
        DATA_DIR / "equipment_catalog_manual_seed.csv",
        equipment_rows,
        ["brand", "name", "category", "recommended_style", "price_band", "data_status", "notes"],
    )

    combo_rows = [
        {
            "combo_name": n,
            "budget_rmb": budget,
            "blade": blade,
            "forehand_rubber": fh,
            "backhand_rubber": bh,
            "target_user": user,
            "evidence_level": level,
            "notes": notes,
        }
        for n, budget, blade, fh, bh, user, level, notes in COMMON_COMBOS
    ]
    write_csv(
        DATA_DIR / "community_combo_seed.csv",
        combo_rows,
        ["combo_name", "budget_rmb", "blade", "forehand_rubber", "backhand_rubber", "target_user", "evidence_level", "notes"],
    )

    question_rows = [
        {
            "question_no": str(no),
            "question": q,
            "options": opts,
            "scoring_axes": axes,
            "status": "MVP_seed",
        }
        for no, q, opts, axes in QUESTIONS
    ]
    write_csv(
        DATA_DIR / "questionnaire_seed.csv",
        question_rows,
        ["question_no", "question", "options", "scoring_axes", "status"],
    )

    watch_rows = [
        {"athlete_name_cn": n, "country": c, "status": s, "next_verification_action": a}
        for n, c, s, a in CN_ATHLETE_WATCHLIST
    ]
    write_csv(
        DATA_DIR / "cn_athlete_watchlist_seed.csv",
        watch_rows,
        ["athlete_name_cn", "country", "status", "next_verification_action"],
    )

    source_rows = [
        {"source_name": n, "source_type": t, "url": u, "use_case": use}
        for n, t, u, use in SOURCES
    ]
    write_csv(DATA_DIR / "source_registry_seed.csv", source_rows, ["source_name", "source_type", "url", "use_case"])

    (DATA_DIR / "schema.sql").write_text(make_schema_sql(), encoding="utf-8")
    (OUTPUT_DIR / "数据采集与制作流程.md").write_text(make_collection_sop(), encoding="utf-8")

    return {
        "athlete": DATA_DIR / "athlete_equipment_evidence_seed.csv",
        "product": DATA_DIR / "butterfly_product_specs_seed.csv",
        "brand": DATA_DIR / "brand_catalog_seed.csv",
        "equipment": DATA_DIR / "equipment_catalog_manual_seed.csv",
        "combo": DATA_DIR / "community_combo_seed.csv",
        "question": DATA_DIR / "questionnaire_seed.csv",
        "watchlist": DATA_DIR / "cn_athlete_watchlist_seed.csv",
        "source": DATA_DIR / "source_registry_seed.csv",
        "schema": DATA_DIR / "schema.sql",
        "sop": OUTPUT_DIR / "数据采集与制作流程.md",
    }


def make_schema_sql() -> str:
    return """-- 乒器通 MVP 数据库草案（PostgreSQL）
-- 原则：事实表必须带 evidence_id 或 source_status；UGC 不直接进入推荐模型。

CREATE TABLE brands (
  id BIGSERIAL PRIMARY KEY,
  name TEXT NOT NULL,
  origin TEXT,
  official_url TEXT,
  created_at TIMESTAMPTZ DEFAULT now()
);

CREATE TABLE equipment_items (
  id BIGSERIAL PRIMARY KEY,
  brand_id BIGINT REFERENCES brands(id),
  name TEXT NOT NULL,
  category TEXT CHECK (category IN ('blade','rubber','ball','accessory')),
  price_min_rmb NUMERIC,
  price_max_rmb NUMERIC,
  status TEXT DEFAULT 'active',
  created_at TIMESTAMPTZ DEFAULT now()
);

CREATE TABLE blade_specs (
  equipment_id BIGINT PRIMARY KEY REFERENCES equipment_items(id),
  structure TEXT,
  thickness_mm NUMERIC,
  weight_g NUMERIC,
  speed_score NUMERIC,
  control_score NUMERIC,
  reaction_property NUMERIC,
  vibration_property NUMERIC
);

CREATE TABLE rubber_specs (
  equipment_id BIGINT PRIMARY KEY REFERENCES equipment_items(id),
  rubber_type TEXT,
  hardness NUMERIC,
  sponge_thickness TEXT,
  speed_score NUMERIC,
  spin_score NUMERIC,
  arc_score NUMERIC,
  tackiness_level TEXT
);

CREATE TABLE evidence_sources (
  id BIGSERIAL PRIMARY KEY,
  source_url TEXT NOT NULL,
  source_name TEXT,
  evidence_level TEXT CHECK (evidence_level IN ('A','B','C','UGC')),
  captured_at TIMESTAMPTZ DEFAULT now(),
  source_date TEXT,
  checksum TEXT,
  reviewer_user_id BIGINT,
  review_status TEXT DEFAULT 'pending'
);

CREATE TABLE athletes (
  id BIGSERIAL PRIMARY KEY,
  name_cn TEXT,
  name_en TEXT,
  country_code TEXT,
  handedness TEXT,
  grip_type TEXT,
  playing_style TEXT,
  current_rank_scope TEXT,
  created_at TIMESTAMPTZ DEFAULT now()
);

CREATE TABLE athlete_equipment_versions (
  id BIGSERIAL PRIMARY KEY,
  athlete_id BIGINT REFERENCES athletes(id),
  equipment_id BIGINT REFERENCES equipment_items(id),
  slot TEXT CHECK (slot IN ('blade','forehand_rubber','backhand_rubber','unknown')),
  valid_from DATE,
  valid_to DATE,
  evidence_id BIGINT REFERENCES evidence_sources(id),
  confidence_score NUMERIC DEFAULT 0,
  notes TEXT
);

CREATE TABLE user_profiles (
  id BIGSERIAL PRIMARY KEY,
  wechat_openid TEXT UNIQUE,
  nickname TEXT,
  age_range TEXT,
  region TEXT,
  kaiqiu_score INTEGER,
  level_stage TEXT CHECK (level_stage IN ('beginner','intermediate','advanced','professional')),
  playing_style TEXT,
  current_blade_id BIGINT REFERENCES equipment_items(id),
  current_fh_rubber_id BIGINT REFERENCES equipment_items(id),
  current_bh_rubber_id BIGINT REFERENCES equipment_items(id),
  privacy_consent_at TIMESTAMPTZ
);

CREATE TABLE recommendation_cases (
  id BIGSERIAL PRIMARY KEY,
  user_id BIGINT REFERENCES user_profiles(id),
  input_json JSONB NOT NULL,
  result_json JSONB NOT NULL,
  algorithm_version TEXT NOT NULL,
  confidence_score NUMERIC,
  created_at TIMESTAMPTZ DEFAULT now()
);

CREATE TABLE personality_tests (
  id BIGSERIAL PRIMARY KEY,
  user_id BIGINT REFERENCES user_profiles(id),
  answers_json JSONB NOT NULL,
  result_label TEXT NOT NULL,
  result_vector JSONB NOT NULL,
  created_at TIMESTAMPTZ DEFAULT now()
);

CREATE TABLE reviews (
  id BIGSERIAL PRIMARY KEY,
  user_id BIGINT REFERENCES user_profiles(id),
  equipment_id BIGINT REFERENCES equipment_items(id),
  score_overall NUMERIC CHECK (score_overall BETWEEN 0 AND 10),
  score_speed NUMERIC,
  score_spin NUMERIC,
  score_control NUMERIC,
  content TEXT,
  moderation_status TEXT DEFAULT 'pending',
  created_at TIMESTAMPTZ DEFAULT now()
);

CREATE TABLE posts (
  id BIGSERIAL PRIMARY KEY,
  user_id BIGINT REFERENCES user_profiles(id),
  title TEXT NOT NULL,
  content TEXT NOT NULL,
  topic_tags TEXT[],
  moderation_status TEXT DEFAULT 'pending',
  created_at TIMESTAMPTZ DEFAULT now()
);

CREATE TABLE comments (
  id BIGSERIAL PRIMARY KEY,
  user_id BIGINT REFERENCES user_profiles(id),
  target_type TEXT CHECK (target_type IN ('review','post','equipment')),
  target_id BIGINT NOT NULL,
  content TEXT NOT NULL,
  moderation_status TEXT DEFAULT 'pending',
  created_at TIMESTAMPTZ DEFAULT now()
);

CREATE TABLE moderation_queue (
  id BIGSERIAL PRIMARY KEY,
  content_type TEXT NOT NULL,
  content_id BIGINT NOT NULL,
  machine_result JSONB,
  human_result TEXT,
  status TEXT DEFAULT 'pending',
  created_at TIMESTAMPTZ DEFAULT now()
);

CREATE TABLE ad_slots (
  id BIGSERIAL PRIMARY KEY,
  slot_key TEXT UNIQUE NOT NULL,
  page_key TEXT NOT NULL,
  max_height_px INTEGER,
  enabled BOOLEAN DEFAULT false,
  notes TEXT
);

CREATE TABLE marketplace_listings (
  id BIGSERIAL PRIMARY KEY,
  user_id BIGINT REFERENCES user_profiles(id),
  equipment_id BIGINT REFERENCES equipment_items(id),
  condition_level TEXT,
  price_rmb NUMERIC,
  description TEXT,
  status TEXT DEFAULT 'reserved_for_future'
);
"""


def make_collection_sop() -> str:
    return f"""# 乒器通数据采集与制作流程

版本：v0.1  
日期：{TODAY}

## 1. 数据分层

- A 级：品牌官网、运动员官方/品牌赞助页、WTT/ITTF 官方页面、授权商品接口。
- B 级：授权经销商、权威媒体采访、比赛高清图像且能识别装备。
- C 级：小红书、贴吧、论坛、淘宝评论等社区样本，只能作为候选和热度参考。
- UGC：用户主动提交，默认不进入推荐模型，必须审核和证据补全。

## 2. 小红书/淘宝处理原则

- 不抓登录后内容，不绕过反爬，不采集个人隐私，不复制正文长内容。
- 淘宝/天猫商品、价格、销量、评价摘要优先走淘宝开放平台/淘宝联盟授权接口。
- 小红书优先走蒲公英/商业合作、人工样本标注或用户授权提交。
- 采集字段只保留商品名、品牌、型号、搭配文本摘要、价格区间、样本链接、采集时间、证据等级。

## 3. 职业运动员装备验证

1. 先用 WTT 排名页确定“顶尖运动员候选名单”。
2. 搜索品牌官网赞助页和产品页；能找到 Equipment as of 日期的记 A 级。
3. 找不到品牌页时，至少需要两类 B 级证据交叉验证：授权媒体/经销商 + 比赛高清图/采访。
4. 每条装备记录必须记录 valid_from、valid_to、source_url、source_date、reviewer。
5. 运动员换拍、换胶皮时新增版本，不覆盖旧记录。

## 4. 推荐模型上线门槛

- MVP 只使用 A/B 级产品参数和人工规则。
- C 级社区搭配只参与“热度提示”，不参与自动推荐排名。
- UGC 要经过文本安全、人工审核、重复检测和证据补全。
"""


def set_run_font(run, name: str = "Microsoft YaHei", size: float | None = None, color: RGBColor | None = None, bold: bool | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_cell_text(cell, text: str, bold: bool = False, size: float = 9.5, color: RGBColor | None = None):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.clear()
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table, color: str = "D9DEE8", size: str = "4"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_table(table, widths: list[float] | None = None):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            if row_idx == 0:
                set_cell_shading(cell, "F2F4F7")
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
            if widths:
                cell.width = Inches(widths[cell_idx])
    table.allow_autofit = False


def add_paragraph(doc, text: str, style: str | None = None, bold: bool = False, size: float = 10.5, color: RGBColor | None = None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    return p


def add_bullets(doc, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.12
        r = p.add_run(item)
        set_run_font(r, size=10.5)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, size=9.5)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=9)
    style_table(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def configure_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12 if style_name != "Heading 1" else 16)
        style.paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.text = "乒器通 PRD | MVP 立项稿"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=8.5, color=RGBColor(100, 100, 100))

    footer = section.footer.paragraphs[0]
    footer.text = "内部讨论稿"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_run_font(run, size=8.5, color=RGBColor(120, 120, 120))


def add_title_page(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("产品需求报告")
    set_run_font(r, size=12, bold=True, color=RGBColor(100, 100, 100))

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("乒器通微信小程序")
    set_run_font(r, size=26, bold=True, color=RGBColor(0, 0, 0))

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("以“专业选型 + 趣味测试 + 器材评分社区”为核心的乒乓球装备决策产品")
    set_run_font(r, size=13, color=RGBColor(55, 55, 55))

    meta = [
        ("版本", "v0.1 MVP 立项稿"),
        ("日期", TODAY),
        ("目标端", "微信小程序；后续预留 Web 管理后台与商城/二手交易"),
        ("核心目标", "先把选型推荐做准，再通过评分社区和趣味测试提高留存"),
        ("数据原则", "官方/品牌/授权来源优先；社区数据只做候选，必须证据分级"),
    ]
    add_table(doc, ["字段", "内容"], [[k, v] for k, v in meta], widths=[1.25, 5.25])

    add_paragraph(
        doc,
        "一句话定位：给乒乓球爱好者一个能看懂、能信任、能复盘的器材选型助手，降低“买错底板和胶皮”的试错成本。",
        bold=True,
        size=11,
        color=RGBColor(31, 58, 95),
    )


def build_report(paths: dict[str, Path]):
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_title_page(doc)

    doc.add_heading("1. 产品定位与目标", level=1)
    add_paragraph(
        doc,
        "乒器通面向业余爱好者、训练型球友、器材发烧友和教练/陪练人群。MVP 的核心不是做大而全社区，而是先解决“预算有限、参数看不懂、搭配靠传闻”的选型痛点。"
    )
    add_bullets(
        doc,
        [
            "主功能：用户输入预算、水平、年龄、打法、当前器材、训练频率、伤病/控重要求后，输出底板、正手胶皮、反手胶皮的可解释推荐。",
            "增长功能：用 15 题左右的乒乓球场景测试生成打法性格标签，并把测试结果反哺选型参数。",
            "社区功能：围绕底板、胶皮、搭配、打法建立评分和文字帖子，短视频和交易能力作为二期预留。",
            "数据原则：职业运动员装备、产品参数、社区搭配必须分证据等级；低证据数据不直接进入推荐主模型。",
        ]
    )

    doc.add_heading("2. 用户与核心场景", level=1)
    add_table(
        doc,
        ["用户类型", "典型问题", "产品解法"],
        [
            ["新手/回坑球友", "不知道该买成品拍还是自配；预算敏感", "用水平和预算限制推荐容错高、重量合理、价格稳定的方案"],
            ["中级进阶球友", "正手想换国套，反手怕太弹；不知道底板外置/内置差异", "用打法、发力、训练频率匹配底板结构和胶皮硬度"],
            ["高级球友/器材党", "想比较相近器材，关注细参数和球友评价", "器材库、评分、评论、同打法用户反馈和职业装备参考"],
            ["教练/陪练", "需要给不同水平学员快速配拍", "提供学员档案、阶段推荐和可复制方案"],
            ["内容贡献者", "愿意分享搭配经验或职业球员装备线索", "UGC 提交、证据上传、审核通过后获得贡献标识"],
        ],
        widths=[1.35, 2.1, 3.05],
    )

    doc.add_heading("3. 功能范围与优先级", level=1)
    add_table(
        doc,
        ["模块", "MVP 必做", "P1 增强", "P2 预留"],
        [
            ["选型推荐", "预算/水平/年龄/打法输入；底板+正反手胶皮组合；推荐理由", "多方案对比、替代型号、价格波动", "基于真实成交和用户反馈的学习排序"],
            ["趣味测试", "15 题场景题；生成打法性格和推荐修正", "分享海报、好友对比", "赛事/训练数据导入后动态更新"],
            ["器材库评分", "底板/胶皮详情、10 分制评分、维度分", "同打法筛选、榜单、避坑标签", "专家点评、视频评测聚合"],
            ["社区", "文字帖、评论、点赞、收藏", "话题、精华、用户等级", "短视频、直播、训练日志"],
            ["用户体系", "微信登录、个人资料、当前装备、地区、开球网积分手填", "历史装备、训练目标", "认证教练/达人"],
            ["审核后台", "评论/帖子机审+人工审核", "装备证据审核、举报处理", "风控、黑名单、申诉"],
            ["商业化", "小广告位预留，不遮挡主流程", "品牌合作页、器材导购", "一手/二手商城、担保交易"],
        ],
        widths=[1.25, 1.9, 1.75, 1.6],
    )

    doc.add_heading("4. 选型推荐系统", level=1)
    add_paragraph(
        doc,
        "推荐逻辑建议先用“规则引擎 + 参数向量 + 专家校准”，不要一开始就追求黑盒模型。用户要知道为什么推荐这套，而不是只看到一个神秘结论。"
    )
    add_table(
        doc,
        ["输入维度", "字段示例", "推荐影响"],
        [
            ["预算", "总预算、是否接受二手、单项预算上限", "限定候选池，避免推荐不可购买方案"],
            ["水平阶段", "初级/中级/高级/专业；训练频率", "控制底板速度、胶皮硬度和容错难度"],
            ["年龄与身体", "年龄段、肩肘腕不适、控重要求", "过滤过重、过硬、需要高发力门槛的配置"],
            ["打法", "近台快攻、两面弧圈、控制、削球、颗粒", "决定底板结构、套胶类型和正反手差异"],
            ["当前装备", "底板、正手、反手、满意/不满意点", "做升级路径，而不是完全重置"],
            ["球感偏好", "持球、通透、弹性、硬度、台内控制", "修正同价位候选的排序"],
            ["场景", "比赛、训练、娱乐、教练推荐", "调整稳定性和上限权重"],
        ],
        widths=[1.3, 2.2, 3.0],
    )
    add_bullets(
        doc,
        [
            "装备向量：速度、旋转、控制、弧线、硬度、弹性、持球、重量、甜区、容错、价格、学习门槛。",
            "用户向量：水平、力量、训练频率、主动性、风险偏好、台内能力、相持能力、控重需求。",
            "输出结构：主推 1 套、备选 2 套、升级路径 1 条、谨慎提示 1 条。",
            "解释原则：每条推荐必须说明“为什么适合你”“可能不适合什么人”“同价位替代型号”。",
        ]
    )

    doc.add_heading("5. 趣味测试设计", level=1)
    add_paragraph(
        doc,
        "趣味测试不只是营销玩法，应当把答案映射成推荐系统可用的打法向量。外显标签可以有梗，后台标签必须中性、可计算。"
    )
    add_table(
        doc,
        ["结果标签", "后台画像", "推荐倾向"],
        [
            ["稳健型快攻打法", "低失误、近台、节奏快", "中等速度、高控制、反手稳定优先"],
            ["暴冲型三板斧", "高主动、高风险、追求前三板", "正手上限和发力反馈优先，提示容错风险"],
            ["算盘型控场弧圈", "旋转变化、线路意识强", "持球、旋转、弧线和台内控制优先"],
            ["老登型磨洋工打法", "低风险、多拍、防守反击", "高控制、防守容错、重量友好，标签可由用户选择是否展示"],
            ["科研型拧拉党", "喜欢研究参数和技术细节", "可提供更细参数解释和相近器材对比"],
            ["防反型铁壁", "防守稳定、等待反击", "防守板/控制型配置，颗粒选项作为分支"],
        ],
        widths=[1.55, 2.05, 2.9],
    )
    add_paragraph(doc, f"首版 15 题已生成到：{paths['question'].name}。", bold=True, color=RGBColor(31, 58, 95))

    doc.add_heading("6. 评分社区与内容审核", level=1)
    add_bullets(
        doc,
        [
            "评分机制：总分 10 分，速度、旋转、控制、性价比、易上手、做工为分项；评分必须绑定用户水平和打法，避免“高阶器材被新手低分误伤”。",
            "评论机制：短评、长评、当前水平、使用时长、搭配底板/胶皮、是否购买凭证可选。",
            "帖子机制：MVP 做文字帖和图片证据，先覆盖选型、搭配、打法讨论、装备线索提交。",
            "审核机制：微信内容安全接口先审文本，命中风险进入人工队列；装备线索类内容必须走证据补全。",
            "风控机制：新用户限频，重复文本/营销外链/引战词进入灰度；评分异常用权重降级而不是直接删除。",
        ]
    )

    doc.add_heading("7. 数据库与证据治理", level=1)
    add_paragraph(
        doc,
        "数据库要同时支持产品参数、职业运动员装备时序、用户推荐、社区评分和未来交易。关键不是表多，而是每条关键事实都能追溯来源。"
    )
    add_table(
        doc,
        ["数据域", "核心表", "注意事项"],
        [
            ["器材", "brands, equipment_items, blade_specs, rubber_specs", "产品参数必须标明来源；不同品牌评分体系不能直接横向等同"],
            ["运动员", "athletes, athlete_equipment_versions, evidence_sources", "支持 valid_from/valid_to，换拍换胶不覆盖旧记录"],
            ["推荐", "recommendation_cases, personality_tests", "保留输入、版本、结果和用户反馈，便于迭代"],
            ["社区", "reviews, posts, comments, moderation_queue", "内容审核状态独立，前台只展示通过内容"],
            ["商业化", "ad_slots, marketplace_listings", "先预留表结构和入口，MVP 不干扰主流程"],
        ],
        widths=[1.25, 2.65, 2.6],
    )
    add_paragraph(doc, f"数据库草案已生成到：{paths['schema'].name}。", bold=True, color=RGBColor(31, 58, 95))

    doc.add_heading("8. 首版数据资产", level=1)
    add_paragraph(
        doc,
        "首版数据分为三类：A 级可入推荐/展示的官方数据、C 级待验证的社区常见搭配、以及中国职业球员待验证清单。这样既能启动产品，又能保护可信度。"
    )
    add_table(
        doc,
        ["文件", "内容", "使用方式"],
        [
            ["brand_catalog_seed.csv", "主流品牌清单、国家/地区、品类、官网", "器材库品牌基础表"],
            ["equipment_catalog_manual_seed.csv", "常见底板/胶皮种子", "候选池，参数补齐前只做展示/人工规则"],
            ["butterfly_product_specs_seed.csv", "Butterfly 官方产品页可解析参数", "可作为 A 级产品参数样例"],
            ["athlete_equipment_evidence_seed.csv", "Butterfly 官方赞助页运动员装备记录", "A 级运动员装备证据表"],
            ["community_combo_seed.csv", "社区常见搭配种子", "仅做热度和采样任务，不直接用于主推荐"],
            ["cn_athlete_watchlist_seed.csv", "中国职业球员待验证清单", "采集任务池，不能当前台事实"],
            ["questionnaire_seed.csv", "趣味测试 15 题", "MVP 测试题库"],
        ],
        widths=[2.15, 2.1, 2.25],
    )

    doc.add_heading("9. 小红书、淘宝等平台采集流程", level=1)
    add_paragraph(
        doc,
        "不建议用爬虫硬抓小红书、淘宝登录后页面。这样不仅稳定性差，也会有平台协议和隐私风险。正确做法是把平台内容当成“候选热度”，通过授权接口、人工样本和用户提交来沉淀。"
    )
    add_bullets(
        doc,
        [
            "淘宝/天猫：通过淘宝开放平台或淘宝联盟接口获取商品标题、价格、店铺、类目、销量/佣金等授权字段；评论只做聚合摘要，不搬运原文。",
            "小红书：通过蒲公英/合作达人/人工标注样本获取搭配趋势；只保存关键词、型号、搭配、链接、采集日期和证据等级。",
            "论坛/社区：公开页面只采摘要和链接，严禁采集手机号、微信号、私信等个人信息。",
            "入库规则：同一搭配需要达到样本数阈值，且至少经过一个人工审核标签，才能从 C 级升级到 B 级候选。",
        ]
    )

    doc.add_heading("10. 微信小程序技术框架", level=1)
    add_table(
        doc,
        ["层级", "建议方案", "原因"],
        [
            ["小程序端", "原生小程序 + TDesign WeChat", "组件稳定、贴近微信生态、后续维护成本低"],
            ["后端", "Node.js/NestJS 或 Java Spring Boot", "推荐规则、审核、社区和未来交易都需要自有服务"],
            ["数据库", "PostgreSQL + Redis", "关系数据清晰，Redis 用于排行榜、热门评分和限流"],
            ["对象存储", "腾讯云 COS", "帖子图片、证据截图、未来短视频素材"],
            ["搜索", "Meilisearch/Elasticsearch", "器材名、别名、帖子和评论检索"],
            ["审核", "微信内容安全接口 + 人工后台", "先机审再人工，保证评论区质量"],
            ["管理后台", "React/Vue Web Admin", "处理装备证据、评论审核、广告位、举报"],
        ],
        widths=[1.2, 2.3, 3.0],
    )

    doc.add_heading("11. 页面结构", level=1)
    add_table(
        doc,
        ["页面", "主要内容", "MVP 要点"],
        [
            ["首页", "选型入口、测试入口、热门器材、轻量广告位", "首屏直接进入选型，广告不超过底部/信息流小位"],
            ["选型问卷", "预算、水平、打法、年龄、当前装备、偏好", "问题少而准，支持跳过不确定项"],
            ["推荐结果", "主推方案、备选方案、理由、风险提示", "能保存到个人资料，能进入器材详情"],
            ["趣味测试", "15 题、进度、结果页", "分享海报和一键带入选型"],
            ["器材库", "品牌、底板、胶皮、评分榜", "筛选条件围绕打法和水平"],
            ["器材详情", "参数、评分、评论、常见搭配、职业参考", "证据等级展示，避免误导"],
            ["社区", "帖子流、话题、发布、评论", "MVP 只做文字和图片证据"],
            ["我的", "个人资料、当前装备、开球网积分、收藏、历史推荐", "资料可手填，明确仅供参考"],
            ["后台", "内容审核、证据审核、器材维护", "不上前台，但必须同步做"],
        ],
        widths=[1.2, 2.35, 2.95],
    )

    doc.add_heading("12. 广告、商城与二手交易预留", level=1)
    add_bullets(
        doc,
        [
            "广告位只放在首页信息流中段、器材详情页底部、推荐结果页底部，不遮挡选型流程。",
            "MVP 广告位默认关闭，只保留 ad_slots 配置和页面容器，避免影响核心体验。",
            "商城预留商品实体、价格快照和跳转字段；一开始可做导购，不急着做自营交易。",
            "二手交易需要额外做身份、担保、纠纷、违禁品和售后流程，建议 P2 后再开。",
        ]
    )

    doc.add_heading("13. 里程碑计划", level=1)
    add_table(
        doc,
        ["阶段", "周期", "交付物"],
        [
            ["第 0 阶段：数据底座", "1-2 周", "品牌/器材/运动员证据表、采集 SOP、审核标准"],
            ["第 1 阶段：原型", "1 周", "小程序页面原型、选型问卷、推荐结果结构"],
            ["第 2 阶段：MVP 开发", "4-6 周", "登录、选型、测试、器材库、评分、评论、后台审核"],
            ["第 3 阶段：内测", "2 周", "50-200 名球友测试，收集推荐准确率和评论质量"],
            ["第 4 阶段：公测", "2-4 周", "补齐数据、优化推荐、开放内容贡献机制"],
            ["第 5 阶段：商业化试点", "公测后", "广告位、导购、品牌合作、二手交易评估"],
        ],
        widths=[1.6, 1.0, 3.9],
    )

    doc.add_heading("14. 成功指标", level=1)
    add_table(
        doc,
        ["指标", "MVP 目标", "说明"],
        [
            ["选型完成率", ">= 55%", "进入问卷后完成并查看结果的比例"],
            ["推荐保存率", ">= 25%", "结果页保存或收藏方案的比例"],
            ["推荐满意度", ">= 4.1/5", "内测问卷或结果页反馈"],
            ["器材详情转化", ">= 35%", "推荐结果点击器材详情"],
            ["有效评论率", ">= 70%", "通过审核且非灌水的评论占比"],
            ["A/B 级数据覆盖", "MVP 核心候选池 >= 60%", "可用于推荐的可信参数覆盖率"],
        ],
        widths=[1.6, 1.3, 3.6],
    )

    doc.add_heading("15. 风险与对策", level=1)
    add_table(
        doc,
        ["风险", "影响", "对策"],
        [
            ["装备数据真假混杂", "推荐失准、用户信任受损", "证据等级、人工审核、来源链接、版本化记录"],
            ["平台采集违规", "数据不可用、合规风险", "只走授权接口/人工样本/用户授权提交"],
            ["评论区营销和吵架", "社区氛围变差", "敏感词、限频、举报、人工审核、账号权重"],
            ["推荐过度承诺", "用户买错后失望", "展示适配理由和风险，不承诺“唯一最优”"],
            ["广告伤害体验", "核心转化下降", "广告默认关闭，小区域、低频、可配置"],
            ["二手交易纠纷", "客服与法律成本", "P2 后再做，先用导购/信息撮合验证需求"],
        ],
        widths=[1.55, 1.8, 3.15],
    )

    doc.add_heading("附录 A：首版官方运动员装备样例", level=1)
    add_table(
        doc,
        ["运动员", "国家/地区", "装备", "类别", "来源日期"],
        [
            ["Fan Zhendong", "CHN", "Fan Zhendong ALC; Dignics 09C", "底板/套胶", "2022-07-19"],
            ["Tomokazu Harimoto", "JPN", "Harimoto Tomokazu Innerforce Super ALC; Zyre 03", "底板/套胶", "2025-10-11"],
            ["Miwa Harimoto", "JPN", "Harimoto Tomokazu Innerforce Super ALC; Dignics 09C; Dignics 05", "底板/套胶", "2025-12-18"],
            ["Lin Yun-Ju", "TPE", "Viscaria Super ALC; Zyre 03; Dignics 05", "底板/套胶", "2026-01-05"],
            ["Patrick Franziska", "GER", "Franziska Innerforce ZLC; Zyre 03; Dignics 09C", "底板/套胶", "2026-06-15"],
            ["Kanak Jha", "USA", "Timo Boll ALC; Zyre 03", "底板/套胶", "2026-02-12"],
        ],
        widths=[1.35, 0.9, 2.95, 0.8, 1.0],
    )
    add_paragraph(doc, "完整记录见 athlete_equipment_evidence_seed.csv；该表按官网赞助页记录，后续仍需补充不同时期装备版本。")

    doc.add_heading("附录 B：参考来源", level=1)
    for name, source_type, url, use_case in SOURCES:
        add_paragraph(doc, f"{name}：{url}（{use_case}）", size=9.5)

    doc.save(REPORT_PATH)


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    paths = write_seed_data()
    build_report(paths)
    print(REPORT_PATH)


if __name__ == "__main__":
    main()
